#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jul  1 13:31:02 2024

@author: ashutoshgoenka
"""

import pandas as pd
import numpy as np
import streamlit as st
from PIL import Image
# from exif import Image as Image2
from PIL.ExifTags import TAGS
import PIL
import os
import zipfile
from zipfile import ZipFile, ZIP_DEFLATED
import pathlib
import shutil
import docx
import docxtpl
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches
import random
from random import randint
from streamlit import session_state



state = session_state
if "key" not in state:
    state["key"] = str(randint(1000, 100000000))
    
if 'counter' not in state: 
    state["counter"]= 0
    
if "loaded" not in state:
    state["loaded"] = False

st.set_page_config(layout="wide")


if state["loaded"] == False:
    try:
        shutil.rmtree("images_comp")
    except:
        pass






def createfile():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    document.save("Image_word_test.docx")


if state["loaded"] == False:
    createfile()
    try:
        os.mkdir("images_comp")
    except:
        pass




def updateTable():
    # global up_files
    createfile()
    global folder
    # global title
    # global selection_selected
    # global df_final
    document = Document("Image_word_test.docx")
    document.add_heading(section_selected + " - Images", 2)
    _, _, files = next(os.walk(folder))
    file_count = len(files)
    st.write(file_count)
    no_of_rows = int(((file_count-1)//3+1)*2)
    table = document.add_table(rows = no_of_rows , cols = 3)
    # st.write("Table Rows " + str(table.rows.))
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Item'     
    # hdr_cells[1].text = 'quantity'
    document.save("Image_word_test.docx")
    counter = 0
    counter_cols = 0
    
    for file in folder.iterdir():
        # name = os.path.splitext(file.name)[0]
        # img_no = int(name.split(" ")[1])
        # adj_img_no = img_no - int(title)
        # st.write(img_no, adj_img_no)
        row_no = (counter//3) *2
        col_no = counter_cols#int(counter_cols - (row_no*3/2))
        # if(row_no>0 and col_no==0):
        #     table.add_row()
        #     table.add_row()
        # st.write(img_no, adj_img_no, row_no, col_no)
        # cell = table.rows[counter].cells[counter_cols]
        cell = table.rows[row_no].cells[col_no]
        cell._element.clear_content()
        picture = cell.add_paragraph().add_run().add_picture('images_comp/'+file.name, width=Inches(2.6))
        cell = table.rows[row_no+1].cells[col_no]
        # cell = table.rows[counter+1].cells[counter_cols]
        # st.write(row_no, col_no)
        cell.text = file.name
        if col_no<2:
            counter_cols = counter_cols + 1
        else:
            # table.add_row()
            counter_cols = 0
        counter = counter+1
    document.add_page_break()
    document.save("Image_word_test.docx")
    
    
    
    
    
# Function to show the images on the webpage
def showImage():
    global up_files
    global total_no_files
    global col1, col2, col3
    
    
    if state["counter"]=="No More Images Left":
        return
        
        
        
    file = up_files[state["counter"]]
    # st.write(state["counter"])
    
    file_exisits = False
    try:
        a = new_width_dict[file.name]
        file_exisits = True
    except:
        pass
    
    if file_exisits == True:
        return
    
    
    im = Image.open(file)
    
    # Displaying Image
    im_width, im_height = im.size 
    original_image_size[file.name] = [im_width, im_height]
    try:
        b = image_size_dict[file.name]
    except:
        image_size_dict[file.name] = [im_width, im_height]
            
    
    
    # st.write(im_width, im_height)
    size_to_scale = min(im_width,im_height)
    # st.write(size_to_scale)
    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
    im_resized = im.crop(box)
    
    try:
        a= next_width_dict[file.name]
        
    except:
        new_width_dict[file.name] = size_to_scale
        new_height_dict[file.name] = size_to_scale
    
    
    st.write(file.name)
    st.write("Image No: " + str(state["counter"] + 1))
    # col1, col2, col3  = st.columns(3)
    with col1:
        col1_img = st.image(file, width=350)
        # oi_width  = st.number_input("width", value = im_width)
        # oi_height = st.number_input("height", value = im_height)
        st.write(im_width, im_height)
        
    
    
    with col2:
        try:
            im_resized = resize_image(im, new_width_dict[file.name], new_height_dict[file.name], 0 , file.name)
        except:
            im_resized = im_resized
            
        
        col2_img = st.image(im_resized, width=350)
        st.write(im_resized.size)
       
        new_width_dict[file.name]  = st.number_input("new width", value = im_resized.size[0], key="nw_"+file.name)
        new_height_dict[file.name] = st.number_input("new height", value = im_resized.size[1], key="nh_"+file.name)
        new_rotation_dict[file.name] =  st.selectbox(
                "Rotation Clockwise",
                # tuple([name] + name_list),
                tuple(rotation_options),
                index= 0,
                # index=name_index_dict[file.name],
                key="nr_"+file.name
                )
        
        # st.write(im_width, im_height)
        st.write(im_resized.size)
        # st.write(im._getexif())
        # st.write(TAGS)
        # exif_data = {TAGS[k]: v for k, v in im._getexif().items() if k in TAGS.keys()}
        # st.write("orientation")
        # st.write(exif_data)
    
    
    with col3:
        try:
            im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name], new_rotation_dict[file.name], file.name )
        except:
            im_resized_final = im_resized
        
        
        col3_img = st.image(im_resized_final, width=350)
        st.write(im_resized_final.size)
    
    
    im_resized_final.save("images_comp/"+file.name)
    
    # state["counter"] += 1
    # if state["counter"] >= total_no_files:
    #     state["counter"]= 0
    # count_file = count_file +1    
    

def button_click_status_update():
    # st.write(state["counter"])
    state["counter"] += 1
    if state["counter"] >= total_no_files:
        state["counter"]= "No More Images Left"


def resize2(img):
    im_width,im_height = img.size
    size_to_scale = min(im_width,im_height)
    box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
    im_resized = im.crop(box)
    return im_resized

def resize(img, new_width):
    width, height  = img.size
    ratio = height/width
    new_height = int(ratio*new_width)
    resized_image = img.resize((new_width, new_height), resample=PIL.Image.LANCZOS)
    
  #   im = Image.open(image_file)
  # im=im.rotate(270, expand=True)
  # im.show()
  # im.save('rotated.jpg')
    return resized_image


def resize_image(img, width, height,  rotation, file_name):

    # exif_data = {TAGS[k]: v for k, v in img._getexif().items() if k in TAGS}
    # st.write("orientation")
    # st.write(exif_data['Orientation'])
    # img_orientation =  exif_data['Orientation']
    # img_rotation_deg =  rotation_degree[rotation]
    # original_orient = rotated_lineup.index(img_orientation)
    # # st.write(rotation_degree, rotation)
    # # st.write(original_orient,img_rotation_deg)
    # new_orient = rotated_lineup[(original_orient + img_rotation_deg)%4]
    # # st.write(original_orient, new_orient)
    
    to_rotate = anti_clockwise_to_clockwise[rotation]
    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    resized_image = img.rotate(to_rotate, expand = True)
    # resized_image.show()
    resized_image.save("images_comp/"+file_name)
    box=((resized_image.width-width)/2,((resized_image.height-height)/2),(resized_image.width+width)/2,((resized_image.height+height)/2))
    im_resized = resized_image.crop(box)
    return im_resized

st.title("Stage 4 - Resize Image")
up_files = st.file_uploader("Upload Image Files", type = ["png", "jpeg", "jpg"] ,accept_multiple_files=True, key=state["key"])
obs_img_list = []
st.write("Upload Observation (Optional)", 4)
obs_file = st.file_uploader("Upload Observation Excel File With Image List Updated", type=['csv','xlsx'],accept_multiple_files=False,key="fileUploader", on_change=createfile)
if obs_file is not None:
    df = pd.read_excel(obs_file)
    df = df.dropna(thresh=5)
    df = df.astype({"Image Number": str})
    # st.write(df)
    obs_img = list(df["Image Number"])
    
    for i in obs_img:
        img_list =  i.split(",")
        for j in img_list:
            obs_img_list.append(j)
    # st.write(obs_img_list)
    
    
    
name_index_dict= {}
image_size_dict = {}
original_image_size= {}
new_width_dict = {}
new_height_dict = {}
new_rotation_dict = {}
section_selected = ""
rotation_options = [0,90,180,270]
rotation_degree = {0: 0, 90: 1, 180: 2, 270: 3}
anti_clockwise_to_clockwise = {0: 0, 90:270, 180:180, 270:90}
rotated_lineup = [1,8,3,6]
col1, col2, col3  = st.columns(3)

# photo_start=1
image_not_used = []
if len(obs_img_list)>0:
    temp_file_list = []
    for file in up_files:
        temp_file_name  = file.name
        used_bool = False
        for i in obs_img_list:
            if i in temp_file_name:
                temp_file_list.append(file)
                used_bool = True
        if used_bool == False:
            image_not_used.append(temp_file_name)
    st.write("Images not displayed: ", image_not_used)
    up_files = temp_file_list

count_file = 0
total_no_files = len(up_files)
# st.write(state["loaded"])
# for file in up_files:
#     st.write(file.name)



if state["loaded"] == False:
    for file in up_files:
        # st.write(file.name)
        im = Image.open(file)
        im2 = resize2(im)
        
        im2.save("images_comp/"+file.name)
        pos = up_files.index(file)
        if pos>= total_no_files-1:
            state["loaded"] = True


if total_no_files >0:
    showImage()
    show_btn = st.button("Show next image", on_click=button_click_status_update)
    
        

    
# for file in up_files:
#     file_exisits = False
#     try:
#         a = new_width_dict[file.name]
#         file_exisits = True
#     except:
#         pass
    
#     if file_exisits == True:
#         continue
        

    
#     # files = os.listdir("images")
#     extensions = ["jpg", "jpeg", "png", "gif", "webp"]
#     im = Image.open(file)
#     ext = file.name.split(".")[-1]
    
    
    
#     # Displaying Image
#     im_width, im_height = im.size 
#     original_image_size[file.name] = [im_width, im_height]
#     try:
#         b = image_size_dict[file.name]
#     except:
#         image_size_dict[file.name] = [im_width, im_height]
        
    
        
    
#     st.write(im_width, im_height)
#     size_to_scale = min(im_width,im_height)
#     st.write(size_to_scale)
#     # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
#     box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
#     im_resized = im.crop(box)
    
#     try:
#         a= next_width_dict[file.name]
        
#     except:
#         new_width_dict[file.name] = size_to_scale
#         new_height_dict[file.name] = size_to_scale
    
    
    # col1, col2, col3  = st.columns(3)
    # with col1:
    #     col1_img = st.image(file, width=350)
    #     # oi_width  = st.number_input("width", value = im_width)
    #     # oi_height = st.number_input("height", value = im_height)
    #     st.write(im_width, im_height)
        
    
    
    # with col2:
    #     try:
    #         im_resized = resize_image(im, new_width_dict[file.name], new_height_dict[file.name], 0 , file.name)
    #     except:
    #         im_resized = im_resized
            
        
    #     col2_img = st.image(im_resized, width=350)
    #     st.write(im_resized.size)
       
    #     new_width_dict[file.name]  = st.number_input("new width", value = im_resized.size[0], key="nw_"+file.name)
    #     new_height_dict[file.name] = st.number_input("new height", value = im_resized.size[1], key="nh_"+file.name)
    #     new_rotation_dict[file.name] =  st.selectbox(
    #             "Rotation Clockwise",
    #             # tuple([name] + name_list),
    #             tuple(rotation_options),
    #             index= 0,
    #             # index=name_index_dict[file.name],
    #             key="nr_"+file.name
    #             )
        
    #     # st.write(im_width, im_height)
    #     st.write(im_resized.size)
    #     # st.write(im._getexif())
    #     # st.write(TAGS)
    #     # exif_data = {TAGS[k]: v for k, v in im._getexif().items() if k in TAGS.keys()}
    #     # st.write("orientation")
    #     # st.write(exif_data)
    
    
    # with col3:
    #     try:
    #         im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name], new_rotation_dict[file.name], file.name )
    #     except:
    #         im_resized_final = im_resized
        
        
    #     col3_img = st.image(im_resized_final, width=350)
    #     st.write(im_resized_final.size)
    
    
    # im_resized_final.save("images_comp/"+file.name)
    # count_file = count_file +1
    
    
    
#  Creating a zip(compressed) folder and add buttons
    
if len(up_files) >0:
    zip_path = "images_compressed.zip"
    directory_to_zip = "images_comp"
    folder = pathlib.Path(directory_to_zip)
    
    with ZipFile(zip_path, 'w', ZIP_DEFLATED) as zip:
        for file in folder.iterdir():
            zip.write(file, arcname=file.name)
            
    
    with open("images_compressed.zip", "rb") as fp:
        btn = st.download_button(
            label="Download ZIP",
            data=fp,
            file_name="images_compressed.zip",
            mime="application/zip"
    )
        
    try:
        with open("Image_word_test.docx", "rb") as fp:
        
            btn_1 = st.button(
                    label="Update Word File",
                    on_click=updateTable,       
                )
    except:
        pass
        # st.write(btn_1)
        
        # if btn_1:
        #     st.write("Running Update Function")
        #     updateTable(up_files)
    
    try:
        with open("Image_word_test.docx", "rb") as fp:
        
            btn_1 = st.download_button(
                    label="Download Word File",
                    data=fp,
                    file_name="Image_word_test",
                    mime="docx"
                    )
    except:
        pass
        
    

    


